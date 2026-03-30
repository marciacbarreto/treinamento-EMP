import streamlit as st
import io
import hashlib
from openai import OpenAI
from audio_recorder_streamlit import audio_recorder
import PyPDF2
import docx

# ------------------------------
# CONFIGURAÇÃO DA PÁGINA
# ------------------------------

st.set_page_config(page_title="Treinamento EMP", layout="wide")

# ------------------------------
# ESTADO DA SESSÃO
# ------------------------------

defaults = {
    "transcricao": "",
    "resposta": "",
    "cv_text": "",
    "last_audio_hash": ""
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ------------------------------
# CLIENTE OPENAI
# ------------------------------

def get_client():
    api_key = st.secrets.get("OPENAI_API_KEY", "")
    return OpenAI(api_key=api_key)

# ------------------------------
# EXTRAIR TEXTO DO CURRÍCULO
# ------------------------------

def extrair_texto_cv(uploaded_file):

    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        pdf = PyPDF2.PdfReader(uploaded_file)
        return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])

    elif name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")

# ------------------------------
# TÍTULO
# ------------------------------

st.title("Treinamento EMP")

# ------------------------------
# EMPRESA
# ------------------------------

empresa = st.text_input("Empresa")

# ------------------------------
# DESCRIÇÃO DA VAGA + CURRÍCULO
# ------------------------------

col1, col2 = st.columns(2)

with col1:

    vaga = st.text_area(
        "Descrição da vaga",
        height=200
    )

with col2:

    uploaded_cv = st.file_uploader(
        "Currículo",
        type=["pdf","docx","txt"]
    )

    if uploaded_cv:
        st.session_state.cv_text = extrair_texto_cv(uploaded_cv)
        st.success("Currículo carregado")

# ------------------------------
# FERRAMENTAS DA EMPRESA
# ------------------------------

st.subheader("Ferramentas que a empresa trabalha (Explicação)")

if empresa and vaga:

    with st.spinner("Identificando ferramentas utilizadas pela empresa..."):

        client = get_client()

        resposta_tools = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """
Identifique as principais ferramentas utilizadas pela empresa
com base na descrição da vaga.

Liste até 6 ferramentas utilizadas nessa área.

Formato obrigatório:
Ferramenta (explicação curta) / Ferramenta (explicação curta)

As ferramentas devem estar relacionadas a:
operações, CRM, análise de dados, automação e gestão de processos.
"""
                },
                {
                    "role": "user",
                    "content": f"""
Empresa: {empresa}

Descrição da vaga:
{vaga}

Currículo:
{st.session_state.cv_text}
"""
                }
            ],
            max_tokens=120
        )

        ferramentas = resposta_tools.choices[0].message.content

        st.info(ferramentas)

# ------------------------------
# BOTÕES
# ------------------------------

colb1, colb2, colb3 = st.columns(3)

with colb1:
    iniciar = st.button("Iniciar")

with colb2:
    atualizar = st.button("Atualizar")

with colb3:
    encerrar = st.button("Encerrar")

st.divider()

# ------------------------------
# PERGUNTA DA ENTREVISTA
# ------------------------------

st.subheader("Pergunta da entrevista")

pergunta_digitada = st.text_input(
    "Digite a pergunta ou use o microfone"
)

audio = audio_recorder(text="Click to record")

client = get_client()

# ------------------------------
# PERGUNTA DIGITADA
# ------------------------------

if pergunta_digitada:

    st.session_state.transcricao = pergunta_digitada

# ------------------------------
# PERGUNTA POR ÁUDIO
# ------------------------------

elif audio:

    audio_hash = hashlib.sha1(audio).hexdigest()

    if audio_hash != st.session_state.last_audio_hash:

        st.session_state.last_audio_hash = audio_hash

        with st.spinner("Transcrevendo pergunta..."):

            audio_file = io.BytesIO(audio)
            audio_file.name = "audio.wav"

            transcricao = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )

            st.session_state.transcricao = transcricao.text

# ------------------------------
# TRANSCRIÇÃO
# ------------------------------

st.subheader("Transcrição da pergunta")

st.text_area(
    "Pergunta detectada",
    value=st.session_state.transcricao,
    height=100
)

# ------------------------------
# GERAR RESPOSTA
# ------------------------------

if st.session_state.transcricao:

    with st.spinner("Gerando resposta estratégica..."):

        resposta = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """
Responda como um humano em entrevista, de forma natural e direta.

Regras:
- Seja rápido e objetivo
- Não invente informação
- Não repita conteúdo
- Use tom coloquial, como conversa
- Só aprofunde se a pergunta exigir
- Se a pergunta for simples, responda em 1 ou 2 frases
- Evite linguagem robótica ou técnica excessiva
"""
                },
                {
                    "role": "user",
                    "content": f"""
Empresa: {empresa}

Descrição da vaga:
{vaga}

Currículo:
{st.session_state.cv_text}

Pergunta:
{st.session_state.transcricao}
"""
                }
            ],
            max_tokens=120
        )

        st.session_state.resposta = resposta.choices[0].message.content

# ------------------------------
# RESPOSTA
# ------------------------------

st.subheader("Resposta estratégica")

st.text_area(
    "Resposta baseada na vaga, currículo e pergunta",
    value=st.session_state.resposta,
    height=180
)


2-------------
import streamlit as st
import io
import hashlib
import re
from openai import OpenAI
from audio_recorder_streamlit import audio_recorder
import PyPDF2
import docx

# ------------------------------
# CONFIGURAÇÃO DA PÁGINA
# ------------------------------

st.set_page_config(page_title="Treinamento EMP", layout="wide")

# ------------------------------
# CSS - FONTE 11
# ------------------------------

st.markdown("""
<style>
div[data-testid="stTextArea"] textarea {
    font-size: 11px !important;
}
</style>
""", unsafe_allow_html=True)

# ------------------------------
# ESTADO DA SESSÃO
# ------------------------------

defaults = {
    "transcricao": "",
    "resposta": "",
    "cv_text": "",
    "last_audio_hash": "",
    "ferramentas_resposta": ""
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ------------------------------
# CLIENTE OPENAI
# ------------------------------

def get_client():
    api_key = st.secrets.get("OPENAI_API_KEY", "")
    return OpenAI(api_key=api_key)

client = get_client()

# ------------------------------
# EXTRAIR TEXTO DO CURRÍCULO
# ------------------------------

def extrair_texto_cv(uploaded_file):
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        pdf = PyPDF2.PdfReader(uploaded_file)
        return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])

    elif name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")

# ------------------------------
# CLASSIFICAÇÃO DA PERGUNTA
# ------------------------------

def tipo_pergunta(texto: str) -> str:
    if not texto:
        return "geral"

    t = texto.lower().strip()

    # Trajetória / apresentação
    if any(p in t for p in [
        "trajetoria", "trajetória", "fale sobre você", "conte sobre você",
        "apresente-se", "me fale da sua carreira", "conte sua carreira"
    ]):
        return "trajetoria"

    # STAR / comportamental
    if any(p in t for p in [
        "conte um case", "me conte um case", "fale de um case",
        "conte uma situação", "me fale de uma situação",
        "desafio", "problema", "erro", "conflito", "resultado",
        "exemplo", "case da gol", "conte um exemplo", "me dê um exemplo",
        "situação", "quando você", "ocasião em que"
    ]):
        return "star"

    # Técnica
    if any(p in t for p in [
        "sql", "power bi", "dashboard", "dados", "query", "processo",
        "como fazer", "como funciona", "fluxo", "arquitetura", "indicador",
        "kpi", "causa raiz"
    ]):
        return "tecnica"

    return "geral"

# ------------------------------
# VALIDAÇÃO STAR
# ------------------------------

def validar_star(texto: str) -> bool:
    if not texto:
        return False

    padroes = [
        r"(?im)^\s*S\s*\(\s*Situa[cç][aã]o\s*\)\s*:",
        r"(?im)^\s*T\s*\(\s*Tarefa\s*\)\s*:",
        r"(?im)^\s*A\s*\(\s*A[cç][aã]o\s*\)\s*:",
        r"(?im)^\s*R\s*\(\s*Resultado\s*\)\s*:"
    ]

    return all(re.search(p, texto) for p in padroes)

# ------------------------------
# TÍTULO
# ------------------------------

st.title("Treinamento EMP")

# ------------------------------
# EMPRESA
# ------------------------------

empresa = st.text_input("Empresa")

# ------------------------------
# DESCRIÇÃO DA VAGA + CURRÍCULO
# ------------------------------

col1, col2 = st.columns(2)

with col1:
    vaga = st.text_area("Descrição da vaga", height=200)

with col2:
    uploaded_cv = st.file_uploader("Currículo", type=["pdf", "docx", "txt"])

    if uploaded_cv:
        st.session_state.cv_text = extrair_texto_cv(uploaded_cv)
        st.success("Currículo carregado")

# ------------------------------
# PERGUNTA DA ENTREVISTA
# ------------------------------

st.subheader("Pergunta da entrevista")

pergunta_digitada = st.text_input("Digite a pergunta ou use o microfone")
audio = audio_recorder(text="Click to record")

# ------------------------------
# PERGUNTA DIGITADA / ÁUDIO
# ------------------------------

if pergunta_digitada:
    st.session_state.transcricao = pergunta_digitada

elif audio:
    audio_hash = hashlib.sha1(audio).hexdigest()

    if audio_hash != st.session_state.last_audio_hash:
        st.session_state.last_audio_hash = audio_hash

        with st.spinner("Transcrevendo pergunta..."):
            audio_file = io.BytesIO(audio)
            audio_file.name = "audio.wav"

            transcricao = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )

            st.session_state.transcricao = transcricao.text

# ------------------------------
# TRANSCRIÇÃO
# ------------------------------

st.subheader("Transcrição da pergunta")

st.text_area(
    "Pergunta detectada",
    value=st.session_state.transcricao,
    height=100
)

# ------------------------------
# GERAR RESPOSTA
# ------------------------------

if st.session_state.transcricao:
    tipo = tipo_pergunta(st.session_state.transcricao)

    if tipo == "trajetoria":
        prompt_extra = """
A pergunta é de trajetória/apresentação.

Responda como narrativa profissional, sem STAR.
Estrutura:
- início da carreira
- evolução
- experiências mais relevantes
- momento atual
- conexão com a vaga

A resposta deve ser fluida, natural e objetiva.
"""

    elif tipo == "star":
        prompt_extra = """
A pergunta é comportamental/case.

Responder obrigatoriamente neste formato visível:

S (Situação):
- Contextualize rapidamente onde e quando aconteceu.

T (Tarefa):
- Explique qual era o problema ou desafio específico.

A (Ação):
- Descreva exatamente o que você fez.
- Foque no "eu", não apenas no "nós".

R (Resultado):
- Mostre o resultado com números, ganhos, melhoria, redução de problema ou aprendizado.

Regras:
- Os blocos S, T, A e R devem aparecer escritos.
- Cada bloco deve ter 1 ou 2 linhas.
- Não responder em texto corrido.
- Não pode faltar nenhuma das quatro partes.
"""

    elif tipo == "tecnica":
        prompt_extra = """
A pergunta é técnica.

Responder de forma direta:
- lógica simples
- passos curtos
- ferramenta usada
- código apenas se necessário
"""

    else:
        prompt_extra = """
Responder de forma clara, curta, natural e conectada à vaga.
"""

    with st.spinner("Gerando resposta estratégica..."):
        resposta = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""
Responda como um humano em entrevista.

Regras gerais:
- Seja claro, direto e natural
- Não invente informação
- Não repita conteúdo
- Use o currículo e a vaga como base
- Seja objetivo

{prompt_extra}
"""
                },
                {
                    "role": "user",
                    "content": f"""
Empresa: {empresa}

Descrição da vaga:
{vaga}

Currículo:
{st.session_state.cv_text}

Pergunta:
{st.session_state.transcricao}
"""
                }
            ],
            max_tokens=260
        )

        resposta_texto = resposta.choices[0].message.content.strip()

        # Reescrita automática se for STAR e não vier no formato correto
        if tipo == "star" and not validar_star(resposta_texto):
            with st.spinner("Ajustando resposta para o formato STAR..."):
                reescrita = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {
                            "role": "system",
                            "content": """
Reescreva a resposta obrigatoriamente em STAR visível:

S (Situação):
...

T (Tarefa):
...

A (Ação):
...

R (Resultado):
...

Regras:
- Não inventar informação
- Cada bloco deve ter 1 ou 2 linhas
- Resposta curta e objetiva
- Manter aderência ao currículo e à pergunta
"""
                        },
                        {
                            "role": "user",
                            "content": f"""
Pergunta:
{st.session_state.transcricao}

Resposta atual:
{resposta_texto}
"""
                        }
                    ],
                    max_tokens=260
                )

                resposta_texto = reescrita.choices[0].message.content.strip()

        st.session_state.resposta = resposta_texto

        # ------------------------------
        # FERRAMENTAS UTILIZADAS NA RESPOSTA
        # ------------------------------

        ferramentas_resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """
Analise a resposta e identifique as ferramentas utilizadas, citadas ou implícitas.

Regras:
- Liste somente ferramentas que façam sentido no contexto
- Explique para que serviram na resposta
- Não inventar ferramentas fora do contexto
- Máximo 5 itens

Formato:
Ferramenta - para que serviu
"""
                },
                {
                    "role": "user",
                    "content": f"""
Resposta:
{st.session_state.resposta}

Vaga:
{vaga}
"""
                }
            ],
            max_tokens=180
        )

        st.session_state.ferramentas_resposta = ferramentas_resp.choices[0].message.content.strip()

# ------------------------------
# RESPOSTA
# ------------------------------

st.subheader("Resposta estratégica")

st.text_area(
    "Resposta baseada na vaga, currículo e pergunta",
    value=st.session_state.resposta,
    height=240
)

# ------------------------------
# QUADRO DE FERRAMENTAS
# ------------------------------

if st.session_state.ferramentas_resposta:
    st.subheader("Ferramentas utilizadas na resposta")

    st.text_area(
        "Ferramentas e para que serviram",
        value=st.session_state.ferramentas_resposta,
        height=140
    )
